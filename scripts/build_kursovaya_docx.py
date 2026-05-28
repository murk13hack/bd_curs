#!/usr/bin/env python3
"""
Сборка KURSOVAYA_BD.md → DOCX по ГОСТ 7.32–2017 (стили Word).

Использование (из корня репозитория):
  python scripts/build_kursovaya_docx.py
  python scripts/build_kursovaya_docx.py --input docs/KURSOVAYA_BD.md --output docs/KURSOVAYA_BD.docx
"""

from __future__ import annotations

import argparse
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DEFAULT_IN = DOCS / "KURSOVAYA_BD.md"
DEFAULT_OUT = DOCS / "KURSOVAYA_BD.docx"
REF_DOC = DOCS / "gost_reference.docx"
PREPARED_MD = DOCS / "_KURSOVAYA_BD_prepared.md"

# ГОСТ 7.32–2017 (типичные параметры курсовой; см. примечание в MD)
MARGIN_LEFT = Cm(2.0)
MARGIN_RIGHT = Cm(1.0)
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
LINE_SPACING = 1.5
FIRST_LINE_INDENT = Cm(1.25)
BLACK = RGBColor(0, 0, 0)


def set_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)
    tc_pr.append(borders)


def configure_paragraph_style(
    style,
    *,
    bold: bool = False,
    italic: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=FONT_SIZE,
    space_before=Pt(0),
    space_after=Pt(0),
    indent_cm: float | None = 1.25,
    line_spacing=LINE_SPACING,
    keep_with_next: bool = False,
    page_break_before: bool = False,
    caps: bool = False,
) -> None:
    font = style.font
    font.name = FONT_NAME
    font.size = size
    font.bold = bold
    font.italic = italic
    font.color.rgb = BLACK
    if caps:
        font.all_caps = True
    pf = style.paragraph_format
    pf.alignment = align
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_with_next
    pf.page_break_before = page_break_before
    if indent_cm is None:
        pf.first_line_indent = None
        pf.left_indent = None
    else:
        pf.first_line_indent = Cm(indent_cm)
    # Times New Roman для восточноевропейской кодировки
    rfonts = style.element.rPr.rFonts if style.element.rPr is not None else None
    if style.element.rPr is None:
        style.element.get_or_add_rPr()
    rpr = style.element.rPr
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), FONT_NAME)
    rpr.append(rfonts)


def build_reference_docx(path: Path) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.left_margin = MARGIN_LEFT
    sec.right_margin = MARGIN_RIGHT
    sec.top_margin = MARGIN_TOP
    sec.bottom_margin = MARGIN_BOTTOM

    normal = doc.styles["Normal"]
    configure_paragraph_style(
        normal,
        bold=False,
        indent_cm=1.25,
        space_after=Pt(0),
    )

    # Встроенные Heading 1–9 переопределяем под иерархию ГОСТ (глава ≠ подпункт)
    configure_paragraph_style(
        doc.styles["Heading 1"],
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent_cm=None,
        space_before=Pt(0),
        space_after=Pt(12),
        page_break_before=False,
        caps=True,
    )
    doc.styles["Heading 1"].name = "Heading 1"

    configure_paragraph_style(
        doc.styles["Heading 2"],
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        indent_cm=None,
        space_before=Pt(18),
        space_after=Pt(12),
        page_break_before=True,
        keep_with_next=True,
    )

    configure_paragraph_style(
        doc.styles["Heading 3"],
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        indent_cm=None,
        space_before=Pt(12),
        space_after=Pt(6),
        page_break_before=False,
        keep_with_next=True,
    )

    configure_paragraph_style(
        doc.styles["Heading 4"],
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        indent_cm=None,
        space_before=Pt(6),
        space_after=Pt(6),
        page_break_before=False,
    )

    configure_paragraph_style(
        doc.styles["Heading 5"],
        bold=False,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent_cm=None,
        space_before=Pt(6),
        space_after=Pt(3),
        page_break_before=False,
    )

    configure_paragraph_style(
        doc.styles["Heading 6"],
        bold=False,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        indent_cm=None,
        space_before=Pt(3),
        space_after=Pt(3),
    )

    for style_name in ("List Paragraph", "Compact", "Block Text"):
        if style_name in doc.styles:
            configure_paragraph_style(
                doc.styles[style_name],
                indent_cm=None,
            )

    quote = doc.styles["Quote"]
    configure_paragraph_style(
        quote,
        italic=True,
        indent_cm=1.0,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=Pt(3),
        space_after=Pt(3),
    )
    quote.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Пользовательские стили (pandoc custom-style=...)
    custom_defs = [
        (
            "GOST Title",
            dict(
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                indent_cm=None,
                space_after=Pt(12),
                caps=True,
            ),
        ),
        (
            "GOST Subtitle",
            dict(
                bold=False,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                indent_cm=None,
                space_after=Pt(18),
            ),
        ),
        (
            "GOST TOC Title",
            dict(
                bold=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                indent_cm=None,
                space_before=Pt(0),
                space_after=Pt(12),
                page_break_before=True,
            ),
        ),
        (
            "GOST Figure Caption",
            dict(
                bold=False,
                italic=False,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                indent_cm=None,
                space_before=Pt(6),
                space_after=Pt(6),
            ),
        ),
        (
            "GOST Placeholder",
            dict(
                bold=False,
                italic=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                indent_cm=None,
                space_before=Pt(6),
                space_after=Pt(6),
            ),
        ),
        (
            "GOST Table Caption",
            dict(
                bold=False,
                italic=False,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                indent_cm=None,
                space_before=Pt(6),
                space_after=Pt(3),
            ),
        ),
        (
            "GOST Code",
            dict(
                bold=False,
                italic=False,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                indent_cm=None,
                size=Pt(11),
                line_spacing=1.0,
            ),
        ),
        (
            "GOST Note",
            dict(
                bold=False,
                italic=True,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                indent_cm=1.0,
                size=Pt(12),
            ),
        ),
    ]
    for name, kwargs in custom_defs:
        st = doc.styles.add_style(name, 1)  # paragraph style
        configure_paragraph_style(st, **kwargs)

    # Таблица с границами — образец для Table style
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    for row in tbl.rows:
        for cell in row.cells:
            set_cell_borders(cell)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE
    doc.add_paragraph("")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def preprocess_markdown(text: str) -> str:
    lines = text.splitlines()
    out: list[str] = []
    in_mermaid = False
    in_code = False
    code_lang = ""

    chapter_re = re.compile(r"^## (\d+\.\s)")
    figure_re = re.compile(r"^#### (Рисунок \d+)")
    subsection_re = re.compile(r"^#### (\d+\.\d+\.\d+\.)")

    i = 0
    while i < len(lines):
        line = lines[i]

        # Убрать редакторское примечание про Markdown
        if "Примечание по оформлению" in line and line.strip().startswith(">"):
            i += 1
            while i < len(lines) and (lines[i].strip().startswith(">") or lines[i].strip() == ""):
                i += 1
            continue

        if line.strip().startswith("```"):
            fence = line.strip()[3:].strip()
            if not in_code:
                in_code = True
                code_lang = fence
                if fence == "mermaid":
                    in_mermaid = True
                    out.append("")
                    out.append('::: {custom-style="GOST Placeholder"}')
                    out.append(
                        "[Здесь вставить рисунок — экспорт из docs/diagrams/, "
                        "см. KURSOVAYA_INSERT_CHECKLIST.md]"
                    )
                    out.append(":::")
                    out.append("")
                else:
                    out.append('```{.GOST-Code}')
                    out.append(line)
            else:
                in_code = False
                if in_mermaid:
                    in_mermaid = False
                else:
                    out.append(line)
            i += 1
            continue

        if in_mermaid:
            i += 1
            continue

        # Заголовки: рисунки — отдельный уровень (не путать с 5.3.1)
        m = figure_re.match(line)
        if m:
            out.append("")
            out.append('::: {custom-style="GOST Figure Caption"}')
            out.append(line.replace("#### ", "", 1).strip())
            out.append(":::")
            i += 1
            continue

        m = subsection_re.match(line)
        if m:
            out.append(f"#### {line[4:].strip()}")
            i += 1
            continue

        # Горизонтальные линии Markdown (---) — не выводить в Word
        if re.match(r"^(-{3,}|\*{3,}|_{3,})\s*$", line.strip()):
            i += 1
            continue

        # Титульный блок и содержание — через fenced div (не в заголовке #)
        if line.strip() == "# ПОЯСНИТЕЛЬНАЯ ЗАПИСКА":
            out.append('::: {custom-style="GOST Title"}')
            out.append("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА")
            out.append(":::")
            i += 1
            continue

        if line.strip() == "## к курсовой работе по дисциплине «Базы данных»":
            out.append('::: {custom-style="GOST Subtitle"}')
            out.append("к курсовой работе по дисциплине «Базы данных»")
            out.append(":::")
            i += 1
            continue

        if line.strip() == "## СОДЕРЖАНИЕ":
            out.append('::: {custom-style="GOST TOC Title"}')
            out.append("СОДЕРЖАНИЕ")
            out.append(":::")
            i += 1
            continue

        # Подписи таблиц **Таблица N** —
        m_tab = re.match(r"^\*\*(Таблица \d+.*?)\*\*", line.strip())
        if m_tab:
            cap = m_tab.group(1).strip()
            out.append("")
            out.append('::: {custom-style="GOST Table Caption"}')
            out.append(cap)
            out.append(":::")
            i += 1
            continue

        # Заглушки рисунков
        if "[Вставить рисунок" in line:
            out.append('::: {custom-style="GOST Placeholder"}')
            out.append(line.strip().lstrip("> ").strip())
            out.append(":::")
            i += 1
            continue

        # Ссылки на якоря → только текст
        line = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)

        out.append(line)
        i += 1

    body = "\n".join(out)
    # Pandoc: ## главы = Heading 2 в reference (стиль «глава» с разрывом страницы)
    yaml = (
        "---\n"
        "title: \"\"\n"
        "lang: ru-RU\n"
        "toc: false\n"
        "numbersections: false\n"
        "---\n\n"
    )
    return yaml + body


def add_page_numbers(docx_path: Path) -> None:
    from docx import Document as Doc
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = Doc(docx_path)
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run()
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        fld = parse_xml(
            f"<w:fldSimple {nsdecls('w')} w:instr=\" PAGE \" />"
        )
        run._r.append(fld)
    doc.save(docx_path)


def _remove_paragraph_border(p_element) -> None:
    p_pr = p_element.get_or_add_pPr()
    for tag in ("w:pBdr",):
        existing = p_pr.find(qn(tag))
        if existing is not None:
            p_pr.remove(existing)


def polish_docx(docx_path: Path) -> None:
    """Чёрный текст, без синих заголовков и без горизонтальных линий."""
    from docx import Document as Doc

    doc = Doc(docx_path)
    custom_style_re = re.compile(r'\{custom-style="[^"]*"\}\s*')

    for p in doc.paragraphs:
        _remove_paragraph_border(p._p)
        if p.text.strip():
            cleaned = custom_style_re.sub("", p.text).strip()
            if cleaned != p.text:
                for run in p.runs:
                    run.text = ""
                if cleaned:
                    r = p.add_run(cleaned)
                    r.font.name = FONT_NAME
                    r.font.color.rgb = BLACK
        for run in p.runs:
            run.font.color.rgb = BLACK
            run.font.name = FONT_NAME

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _remove_paragraph_border(p._p)
                    for run in p.runs:
                        run.font.color.rgb = BLACK

    for style in doc.styles:
        try:
            if hasattr(style, "font") and style.font is not None:
                style.font.color.rgb = BLACK
        except (AttributeError, ValueError):
            pass

    doc.save(docx_path)


def patch_tables_borders(docx_path: Path) -> None:
    from docx import Document as Doc

    doc = Doc(docx_path)
    for table in doc.tables:
        table.style = "Table Grid"
        for row in table.rows:
            for cell in row.cells:
                set_cell_borders(cell)
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.first_line_indent = None
                    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = FONT_NAME
                        run.font.size = FONT_SIZE
                        run.font.color.rgb = BLACK
    doc.save(docx_path)


def run_pandoc(md_path: Path, out_path: Path, ref_path: Path) -> None:
    cmd = [
        "pandoc",
        str(md_path),
        "-o",
        str(out_path),
        "--from",
        "markdown+smart+fenced_divs+raw_attribute",
        "--reference-doc",
        str(ref_path),
        "--standalone",
        "--wrap=none",
        "-V",
        "lang=ru-RU",
    ]
    subprocess.run(cmd, check=True, cwd=ROOT)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build GOST DOCX for kursovaya")
    parser.add_argument("--input", type=Path, default=DEFAULT_IN)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUT)
    parser.add_argument("--reference", type=Path, default=REF_DOC)
    args = parser.parse_args()

    if not args.input.exists():
        print(f"Not found: {args.input}", file=sys.stderr)
        return 1

    print("Building GOST reference styles…")
    build_reference_docx(args.reference)

    print("Preparing markdown…")
    raw = args.input.read_text(encoding="utf-8-sig")
    prepared = preprocess_markdown(raw)
    PREPARED_MD.write_text(prepared, encoding="utf-8")

    print("Running pandoc…")
    run_pandoc(PREPARED_MD, args.output, args.reference)

    print("Patching tables and typography…")
    patch_tables_borders(args.output)
    polish_docx(args.output)

    print("Adding page numbers (footer, center)…")
    add_page_numbers(args.output)

    print(f"Done: {args.output}")
    print(f"Reference: {args.reference}")
    print(f"Prepared MD: {PREPARED_MD}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
